from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DARK   = RGBColor(0x1E, 0x29, 0x3B)
GRAY   = RGBColor(0x47, 0x55, 0x69)
LIGHT  = RGBColor(0x94, 0xA3, 0xB8)
ACCENT = RGBColor(0x08, 0x5E, 0x7A)
TBLHDR = 'F1F5F9'


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.size      = Pt(8.5)
    run.bold           = True
    run.font.color.rgb = ACCENT
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bot)
    pPr.append(pBdr)


def body(text, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = GRAY
    return p


def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = 'Table Grid'

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, TBLHDR)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Pt(4)
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = DARK

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Pt(4)
            if isinstance(val, tuple):
                r = p.add_run(val[0])
                r.bold = val[1]; r.font.size = Pt(9.5)
                r.font.name = 'Courier New' if val[1] else None
                r.font.color.rgb = DARK if val[1] else GRAY
            else:
                r = p.add_run(str(val))
                r.font.size = Pt(9.5); r.font.color.rgb = GRAY

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def code_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name      = 'Courier New'
    r.font.size      = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0x5C)


def title_block(subtitle, page_label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('NeedLink')
    r1.bold = True; r1.font.size = Pt(22); r1.font.color.rgb = DARK
    r2 = p.add_run('  ·  ' + subtitle)
    r2.font.size = Pt(13); r2.font.color.rgb = LIGHT

    pm = doc.add_paragraph()
    pm.paragraph_format.space_after = Pt(1)
    rm = pm.add_run('Version 1.0   ·   May 2026   ·   ' + page_label)
    rm.font.size = Pt(9); rm.font.color.rgb = LIGHT; rm.italic = True

    pd = doc.add_paragraph()
    pd.paragraph_format.space_before = Pt(3)
    pd.paragraph_format.space_after  = Pt(10)
    pPr  = pd._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '0D0D0D')
    pBdr.append(bot); pPr.append(pBdr)


# ================================================================
# PAGE 1
# ================================================================

title_block('Technical Documentation', 'Page 1 of 2')

h2('1. Problem and Solution')
body(
    'NGOs managing in-kind donation campaigns typically rely on WhatsApp groups, email '
    'threads, and spreadsheets to communicate what they need, track incoming pledges, and '
    'confirm deliveries. This leads to duplicate pledges, missed deadlines, and no reliable '
    'audit trail. NeedLink solves this by providing a single platform where NGOs post '
    'donation requests and donors discover, pledge, and track those donations through to '
    'delivery. Every step is recorded in a central database with access controls so each '
    'party only sees what is relevant to them.'
)

h2('2. System Architecture')
add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        [('Mobile App',  True), 'Flutter 3 / Dart 3',                'Primary client used daily by donors and NGO staff for all donation workflows.'],
        [('Web Portal',  True), 'React 19 / TypeScript / Vite',      'Website used only for NGO registration and account setup.'],
        [('Backend',     True), 'Supabase (PostgreSQL)',              'Provides the database, user authentication, REST API, and row-level security. All communication is over HTTPS with JWT tokens.'],
    ],
    col_widths=[1.3, 2.1, 3.1]
)

h2('3. Technology Stack')
add_table(
    ['Package', 'Version', 'Role'],
    [
        [('flutter / dart',          True), '^3.11.0',  'UI framework and language for the mobile app'],
        [('supabase_flutter',        True), '^2.12.4',  'Supabase client for authentication and database queries'],
        [('hooks_riverpod',          True), '^3.3.1',   'State management using FutureProvider and StreamProvider'],
        [('go_router',               True), '^17.2.3',  'Declarative routing with ShellRoute tab shells'],
        [('google_fonts',            True), '^6.2.1',   'Typography using Plus Jakarta Sans'],
        [('cached_network_image',    True), '^3.4.1',   'Efficient network image loading and caching'],
        [('react',                   True), '^19.2.5',  'UI framework for the web portal'],
        [('@supabase/supabase-js',   True), '^2.105.4', 'Supabase client for the web portal'],
        [('tailwindcss',             True), '^4.3.0',   'Utility-first CSS styling'],
        [('zod',                     True), '^4.4.3',   'Schema validation for registration forms'],
        [('recharts',                True), '^3.8.1',   'Charts and data visualisation'],
    ],
    col_widths=[2.1, 1.1, 3.3]
)

h2('4. Mobile App Architecture')
add_table(
    ['File / Layer', 'Description'],
    [
        [('main.dart',               True), 'Entry point. Initialises the Supabase client, wraps the app in a Riverpod ProviderScope, and attaches the GoRouter instance.'],
        [('models.dart',             True), 'Defines four typed classes: Profile, Ngo, DonationNeed, and Pledge. Each has a fromJson constructor. DonationNeed exposes computed helpers: progress, remaining, isUrgent, and isOpen.'],
        [('providers.dart',          True), 'Six Riverpod FutureProviders, each wrapping one Supabase query. Widgets never call Supabase directly; they observe a provider and receive loading, data, or error state.'],
        [('router.dart',             True), 'GoRouter config with two ShellRoutes (DonorShell, NgoShell), a splash resolver that reads user role, and a global redirect guard for unauthenticated access.'],
        [('theme.dart',              True), 'Named colour constants (kPrimary, kAccent, kUrgent, kMatched, etc.) and a global ThemeData. No widget hardcodes a colour directly.'],
        [('widgets/app_shell.dart',  True), 'DonorShell and NgoShell — persistent NavigationBar wrappers shared by all tab screens.'],
        [('widgets/need_card.dart',  True), 'Reusable donation card with a gradient hero area, dot pattern painter, urgency badge, and progress bar.'],
        [('screens/auth/',           True), 'LoginScreen and RegisterScreen.'],
        [('screens/donor/',          True), 'DonorHomeScreen, NeedDetailScreen, MyPledgesScreen, TrackingScreen, TrackingDetailScreen, DonorProfileScreen.'],
        [('screens/ngo/',            True), 'NgoHomeScreen, CreateNeedScreen, NgoPledgesScreen, ImpactReportsScreen, NgoSettingsScreen.'],
    ],
    col_widths=[2.0, 4.5]
)


# ================================================================
# PAGE 2
# ================================================================

doc.add_page_break()

title_block('Technical Documentation', 'Page 2 of 2')

# 5. Screens and Features — TABLE
h2('5. Screens and Features')
add_table(
    ['Screen', 'Role', 'Description'],
    [
        [('Donor Home',          True), 'Donor',     'Discovery feed of open needs. Pinned search bar and category filter chips. Pull-to-refresh re-fetches data.'],
        [('Need Detail',         True), 'Donor',     'Full need information with progress bar. A bottom sheet slides up to collect pledge quantity, delivery date, and notes.'],
        [('My Pledges',          True), 'Donor',     'All pledges with summary stats at the top. Filter chips narrow by status. Each entry links to tracking detail.'],
        [('Tracking Detail',     True), 'Donor',     'Delivery timeline for a single pledge, showing each status change from pending to confirmed.'],
        [('Donor Profile',       True), 'Donor',     'Contribution stats, donor tier badge, recent donations list, and account settings including sign-out.'],
        [('NGO Dashboard',       True), 'NGO Admin', 'Overview of total pledges, items received, active needs with progress bars, and incoming pledges to confirm or reject.'],
        [('Create Need',         True), 'NGO Admin', '3-step form: item details and urgency, then quantity and deadline, then a review summary before publishing.'],
        [('NGO Pledges',         True), 'NGO Admin', 'Full pledge list across all needs. Filter by status. Each card has Confirm and Reject buttons with loading states.'],
        [('Impact Reports',      True), 'NGO Admin', 'Key metrics, category breakdown with progress bars, success story cards, and downloadable period report tiles.'],
        [('NGO Settings',        True), 'NGO Admin', 'Organisation profile management and account settings for the NGO admin.'],
    ],
    col_widths=[1.5, 1.0, 4.0]
)

# 6. Key Implementation Patterns — prose
h2('6. Key Implementation Patterns')
body(
    'Supabase supports relational joins within a single API call using an embedded resource '
    'syntax. Fetching a pledge together with its related donation need and the donor profile '
    'in one round trip is written as:'
)
code_line(".select('*, donation_need:donation_needs!inner(*, ngo_id), donor:profiles!donor_id(full_name, phone)')")
body(
    'Because Supabase returns raw JSON as an untyped list, every provider applies an '
    'explicit cast before calling the model fromJson constructor:'
)
code_line("(data as List).map((e) => Model.fromJson(e as Map<String, dynamic>)).toList()")
body(
    'The NeedCard widget uses a CustomPainter to draw a dot grid over a gradient hero area. '
    'Each category maps to its own gradient so cards are visually distinct without text '
    'labels. The CreateNeedScreen manages step state locally using an IndexedStack and wraps '
    'the final submit in a try/catch with a loading indicator and error snackbar on failure.'
)

# 7. Authentication Flow — TABLE
h2('7. Authentication Flow')
add_table(
    ['Step', 'Where', 'What Happens'],
    [
        ['1', 'Register screen',       'User submits name, email, password, and role. App calls Supabase sign-up with role and name passed in the user metadata payload.'],
        ['2', 'Supabase Auth',          'Supabase creates a new row in auth.users and issues a signed JWT session token.'],
        ['3', 'PostgreSQL trigger',     'The on_auth_user_created trigger fires immediately, reads name and role from the metadata, and inserts a matching row in public.profiles.'],
        ['4', 'Splash screen',          'On next launch the app checks for an active session. If none exists the user goes to the login screen.'],
        ['5', 'Role resolution',        'If a session exists the app fetches profiles.role for the current user and routes to the donor shell or the NGO shell accordingly.'],
        ['6', 'Token refresh',          'The Supabase Flutter SDK refreshes the JWT silently in the background. The app never handles tokens manually.'],
    ],
    col_widths=[0.4, 1.6, 4.5]
)

# 8. Data Security — prose
h2('8. Data Security')
body(
    'Row-level security is enabled on every table in the database. Each table has policies '
    'that define exactly which rows a given authenticated user can read, insert, update, or '
    'delete. These policies are enforced inside PostgreSQL itself, not in application code, '
    'which means they cannot be bypassed even if someone sends a crafted request directly '
    'to the REST endpoint. Donors can only see and write their own pledge rows. NGO admins '
    'can only manage needs and pledges that belong to their own organisation. Profiles are '
    'readable by any logged-in user but writable only by the account owner. NGO records and '
    'donation needs are publicly readable so unauthenticated visitors on the web portal can '
    'browse without an account.'
)

doc.save('/home/nzabanita/AndroidStudioProjects/NeedLink/docs/needlink-overview.docx')
print('Done.')
